
from datetime import datetime, timedelta, timezone
from io import BytesIO
from pathlib import Path
from typing import Any, Dict
from urllib.request import Request, urlopen

from docx.shared import Inches

from docx import Document


TEMPLATE_FILENAME = "LuxuryEvermore New Receipt Template Auto.docx"
IMAGE_PLACEHOLDER = "{{Image}}"
# Current image width was 2.6 inches; user requested half size.
IMAGE_WIDTH_INCHES = 1


def _dt8_date_str(dt: Any) -> str:
    """
    Return YYYY-MM-DD in UTC+8.
    Accepts datetime / ISO string / other values (best-effort).
    """
    if not dt:
        return ""
    tz8 = timezone(timedelta(hours=8))
    if isinstance(dt, datetime):
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=timezone.utc)
        return dt.astimezone(tz8).strftime("%Y-%m-%d")
    if isinstance(dt, str):
        # try iso, fallback to first 10 chars
        try:
            d2 = datetime.fromisoformat(dt.replace("Z", "+00:00"))
            if d2.tzinfo is None:
                d2 = d2.replace(tzinfo=timezone.utc)
            return d2.astimezone(tz8).strftime("%Y-%m-%d")
        except Exception:
            return dt[:10]
    return str(dt)[:10]


def replace_placeholder_in_paragraph(paragraph, placeholder: str, replacement: str) -> None:
    """
    Replace placeholder in a paragraph. Supports placeholder spanning multiple runs.
    Keeps formatting as much as possible (based on the placeholder start run).

    This logic is extracted from receipt_template/test.ipynb.
    """
    runs = paragraph.runs
    if not runs:
        return

    full = ""
    run_spans = []  # (start, end)
    for r in runs:
        start = len(full)
        full += r.text or ""
        end = len(full)
        run_spans.append((start, end))

    if placeholder not in full:
        return

    while True:
        idx = full.find(placeholder)
        if idx == -1:
            break

        start_idx = idx
        end_idx = idx + len(placeholder)

        start_run = None
        end_run = None
        for i, (s, e) in enumerate(run_spans):
            if start_run is None and s <= start_idx < e:
                start_run = i
            if s < end_idx <= e:
                end_run = i
                break

        if start_run is None or end_run is None:
            break

        s_run_start, _ = run_spans[start_run]
        e_run_start, _ = run_spans[end_run]
        s_off = start_idx - s_run_start
        e_off = end_idx - e_run_start

        s_text = runs[start_run].text or ""
        e_text = runs[end_run].text or ""

        prefix = s_text[:s_off]
        if start_run == end_run:
            suffix = s_text[e_off:]
            runs[start_run].text = prefix + str(replacement) + suffix
        else:
            runs[start_run].text = prefix + str(replacement)
            for k in range(start_run + 1, end_run):
                runs[k].text = ""
            suffix = e_text[e_off:]
            runs[end_run].text = suffix

        # recompute spans after replacement
        full = ""
        run_spans = []
        for r in runs:
            st = len(full)
            full += r.text or ""
            en = len(full)
            run_spans.append((st, en))


def _replace_in_paragraph_all(paragraph, mapping: Dict[str, Any]) -> None:
    for k, v in mapping.items():
        replace_placeholder_in_paragraph(paragraph, k, "" if v is None else str(v))


def _replace_in_table(table, mapping: Dict[str, Any]) -> None:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                _replace_in_paragraph_all(p, mapping)
            for t in cell.tables:
                _replace_in_table(t, mapping)


def replace_everywhere(doc: Document, mapping: Dict[str, Any]) -> None:
    # body paragraphs
    for p in doc.paragraphs:
        _replace_in_paragraph_all(p, mapping)

    # body tables
    for table in doc.tables:
        _replace_in_table(table, mapping)

    # headers/footers
    for section in doc.sections:
        header = section.header
        footer = section.footer

        for p in header.paragraphs:
            _replace_in_paragraph_all(p, mapping)
        for table in header.tables:
            _replace_in_table(table, mapping)

        for p in footer.paragraphs:
            _replace_in_paragraph_all(p, mapping)
        for table in footer.tables:
            _replace_in_table(table, mapping)


def _fetch_url_bytes(url: str, *, timeout_sec: int = 15) -> BytesIO:
    req = Request(
        url,
        headers={
            "User-Agent": "Mozilla/5.0 (receipt-generator)",
            "Accept": "image/*,*/*;q=0.8",
        },
        method="GET",
    )
    with urlopen(req, timeout=timeout_sec) as resp:
        data = resp.read()
    buf = BytesIO(data)
    buf.seek(0)
    return buf


def _insert_run_after(paragraph, run_idx: int):
    """
    Insert a new run after paragraph.runs[run_idx] and return it.
    """
    runs = paragraph.runs
    if not runs:
        return paragraph.add_run()

    anchor = runs[run_idx]
    new_run = paragraph.add_run()
    # move the new run right after anchor in the underlying XML
    anchor._r.addnext(new_run._r)  # noqa: SLF001 (python-docx internal)
    return new_run


def replace_placeholder_with_image_in_paragraph(paragraph, placeholder: str, image_stream: BytesIO) -> None:
    """
    Replace placeholder with an embedded image in a paragraph (supports spanning runs).
    """
    runs = paragraph.runs
    if not runs:
        return

    # build full text and spans
    full = ""
    run_spans = []
    for r in runs:
        start = len(full)
        full += r.text or ""
        end = len(full)
        run_spans.append((start, end))

    if placeholder not in full:
        return

    while True:
        idx = full.find(placeholder)
        if idx == -1:
            break

        start_idx = idx
        end_idx = idx + len(placeholder)

        start_run = None
        end_run = None
        for i, (s, e) in enumerate(run_spans):
            if start_run is None and s <= start_idx < e:
                start_run = i
            if s < end_idx <= e:
                end_run = i
                break

        if start_run is None or end_run is None:
            break

        s_run_start, _ = run_spans[start_run]
        e_run_start, _ = run_spans[end_run]
        s_off = start_idx - s_run_start
        e_off = end_idx - e_run_start

        s_text = runs[start_run].text or ""
        e_text = runs[end_run].text or ""

        prefix = s_text[:s_off]
        if start_run == end_run:
            suffix = s_text[e_off:]
            runs[start_run].text = prefix + suffix
        else:
            runs[start_run].text = prefix
            for k in range(start_run + 1, end_run):
                runs[k].text = ""
            suffix = e_text[e_off:]
            runs[end_run].text = suffix

        # insert image run after start_run
        image_stream.seek(0)
        r = _insert_run_after(paragraph, start_run)
        r.add_picture(image_stream, width=Inches(IMAGE_WIDTH_INCHES))

        # recompute spans after replacement
        runs = paragraph.runs
        full = ""
        run_spans = []
        for rr in runs:
            st = len(full)
            full += rr.text or ""
            en = len(full)
            run_spans.append((st, en))


def _replace_image_in_table(table, placeholder: str, image_stream: BytesIO) -> None:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                replace_placeholder_with_image_in_paragraph(p, placeholder, image_stream)
            for t in cell.tables:
                _replace_image_in_table(t, placeholder, image_stream)


def replace_image_everywhere(doc: Document, placeholder: str, image_stream: BytesIO) -> None:
    for p in doc.paragraphs:
        replace_placeholder_with_image_in_paragraph(p, placeholder, image_stream)

    for table in doc.tables:
        _replace_image_in_table(table, placeholder, image_stream)

    for section in doc.sections:
        header = section.header
        footer = section.footer

        for p in header.paragraphs:
            replace_placeholder_with_image_in_paragraph(p, placeholder, image_stream)
        for table in header.tables:
            _replace_image_in_table(table, placeholder, image_stream)

        for p in footer.paragraphs:
            replace_placeholder_with_image_in_paragraph(p, placeholder, image_stream)
        for table in footer.tables:
            _replace_image_in_table(table, placeholder, image_stream)


def template_path() -> Path:
    return Path(__file__).resolve().parent / TEMPLATE_FILENAME


def build_receipt_mapping(*, item: Dict[str, Any], sold: Dict[str, Any]) -> Dict[str, Any]:
    """
    Map MongoDB fields -> template placeholders.
    Placeholders are based on receipt_template/test.ipynb.
    """
    # IMPORTANT: do not use internal item.name on receipts.
    # Prefer Shopify title if linked; otherwise leave blank.
    item_name = ""
    try:
        if item.get("shopify_sku_exist"):
            sd = item.get("shopify_details") or {}
            item_name = (sd.get("title") or "").strip()
    except Exception:
        item_name = ""

    sold_price = sold.get("sold_price")
    sold_currency = (sold.get("sold_currency") or "SGD").strip().upper()

    return {
        "{{Buyer}}": sold.get("buyer") or "",
        "{{Receipt_no}}": sold.get("receipt_no") or "",
        "{{Date}}": _dt8_date_str(sold.get("sold_at")),
        "{{Payment_method}}": sold.get("payment_method") or "",
        "{{Payment_status}}": "Paid",
        "{{Additional_notes}}": sold.get("sale_note") or "N.A.",
        "{{Item_name}}": item_name,
        "{{Inclusions}}": sold.get("package_inclusion") or "",
        "{{Amount}}": "" if sold_price is None else str(sold_price),
        "{{Currency}}": sold_currency,
        "{{Total_amount}}": "" if sold_price is None else str(sold_price),
    }


def generate_receipt_docx_bytes(*, item: Dict[str, Any], sold: Dict[str, Any]) -> BytesIO:
    """
    Generate a filled receipt docx into an in-memory BytesIO.
    """
    tpl = template_path()
    doc = Document(str(tpl))
    mapping = build_receipt_mapping(item=item, sold=sold)
    replace_everywhere(doc, mapping)

    # Embed Shopify featured image at {{Image}} if present; otherwise remove placeholder.
    img_url = ""
    try:
        sd = item.get("shopify_details") or {}
        img_url = (sd.get("featured_image") or "").strip()
    except Exception:
        img_url = ""

    if img_url:
        try:
            img_stream = _fetch_url_bytes(img_url)
            replace_image_everywhere(doc, IMAGE_PLACEHOLDER, img_stream)
        except Exception:
            # If fetch/embed fails, just clear the placeholder text.
            replace_everywhere(doc, {IMAGE_PLACEHOLDER: ""})
    else:
        replace_everywhere(doc, {IMAGE_PLACEHOLDER: ""})

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


